from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT_HEX = "1B4F72"          # deep professional blue - name, headings
RULE_HEX = "7FA6C2"            # lighter tint of accent - divider lines
MUTED_HEX = "555555"           # muted gray - meta text (dates, locations)

ACCENT_COLOR = RGBColor.from_string(ACCENT_HEX)
MUTED_COLOR = RGBColor.from_string(MUTED_HEX)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(3)
    border_p = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), RULE_HEX)
    pbdr.append(bottom)
    border_p.append(pbdr)


def _meta_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED_COLOR
    run.font.size = Pt(9.5)


def render_docx(profile: dict, out_path: Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Pt(34)
        section.bottom_margin = Pt(34)
        section.left_margin = Pt(46)
        section.right_margin = Pt(46)

    contact = profile.get("contact", {})

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(contact.get("name", "").strip() or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ACCENT_COLOR

    target_roles = profile.get("target_roles") or []
    if target_roles:
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_after = Pt(6)
        role_run = role_p.add_run(" | ".join(target_roles[:2]))
        role_run.font.size = Pt(12)
        role_run.font.color.rgb = RGBColor.from_string("333333")

    contact_bits = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("location", ""),
        contact.get("linkedin", ""),
        *contact.get("other_links", []),
    ]
    contact_line = " • ".join(b for b in contact_bits if b)
    if contact_line:
        contact_p = doc.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(4)
        run = contact_p.add_run(contact_line)
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED_COLOR

    # header divider rule
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(0)
    rule_pPr = rule_p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pbdr.append(bottom)
    rule_pPr.append(pbdr)

    if profile.get("summary"):
        _add_heading(doc, "Summary")
        doc.add_paragraph(profile["summary"])

    if profile.get("skills"):
        _add_heading(doc, "Skills")
        doc.add_paragraph(" • ".join(profile["skills"]))

    if profile.get("experience"):
        _add_heading(doc, "Experience")
        for exp in profile["experience"]:
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            title_run = line.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            dates = " – ".join(x for x in (exp.get("start", ""), exp.get("end", "")) if x)
            meta_bits = [b for b in (exp.get("location", ""), dates) if b]
            if meta_bits:
                _meta_run(line, "    " + " | ".join(meta_bits))
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)

    if profile.get("projects"):
        _add_heading(doc, "Projects")
        for proj in profile["projects"]:
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            name_run = line.add_run(proj.get("name", ""))
            name_run.bold = True
            if proj.get("link"):
                _meta_run(line, f"  ({proj['link']})")
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            for bullet in proj.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)

    if profile.get("education"):
        _add_heading(doc, "Education")
        for edu in profile["education"]:
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            deg_run = line.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
            deg_run.bold = True
            dates = " – ".join(x for x in (edu.get("start", ""), edu.get("end", "")) if x)
            meta_bits = [b for b in (edu.get("location", ""), dates) if b]
            if meta_bits:
                _meta_run(line, "    " + " | ".join(meta_bits))
            if edu.get("details"):
                doc.add_paragraph(edu["details"])

    if profile.get("publications"):
        _add_heading(doc, "Publications")
        for pub in profile["publications"]:
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            title_run = line.add_run(pub.get("title", ""))
            title_run.bold = True
            if pub.get("co_authored"):
                _meta_run(line, " (co-authored)")
            meta_bits = [b for b in (pub.get("venue", ""), pub.get("date", "")) if b]
            if meta_bits:
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_after = Pt(2)
                _meta_run(meta_p, " | ".join(meta_bits))
            if pub.get("details"):
                doc.add_paragraph(pub["details"])

    if profile.get("certifications"):
        _add_heading(doc, "Certifications")
        doc.add_paragraph(" • ".join(profile["certifications"]))

    if profile.get("languages_spoken"):
        _add_heading(doc, "Languages")
        bits = [
            f"{lang.get('language', '')} ({lang.get('level', '')})" if lang.get("level") else lang.get("language", "")
            for lang in profile["languages_spoken"]
            if lang.get("language")
        ]
        doc.add_paragraph(" • ".join(bits))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def render_txt(profile: dict, out_path: Path) -> Path:
    """A plain-text rendering for ATS parsers that struggle with .docx layout."""
    lines: list[str] = []
    contact = profile.get("contact", {})
    lines.append(contact.get("name", ""))
    lines.append(
        " | ".join(
            b
            for b in (
                contact.get("email", ""),
                contact.get("phone", ""),
                contact.get("location", ""),
                contact.get("linkedin", ""),
                *contact.get("other_links", []),
            )
            if b
        )
    )
    lines.append("")

    if profile.get("summary"):
        lines += ["SUMMARY", profile["summary"], ""]

    if profile.get("skills"):
        lines += ["SKILLS", ", ".join(profile["skills"]), ""]

    if profile.get("experience"):
        lines.append("EXPERIENCE")
        for exp in profile["experience"]:
            dates = " - ".join(x for x in (exp.get("start", ""), exp.get("end", "")) if x)
            lines.append(f"{exp.get('title', '')} — {exp.get('company', '')} ({dates})")
            for bullet in exp.get("bullets", []):
                lines.append(f"  - {bullet}")
        lines.append("")

    if profile.get("projects"):
        lines.append("PROJECTS")
        for proj in profile["projects"]:
            lines.append(proj.get("name", ""))
            for bullet in proj.get("bullets", []):
                lines.append(f"  - {bullet}")
        lines.append("")

    if profile.get("education"):
        lines.append("EDUCATION")
        for edu in profile["education"]:
            dates = " - ".join(x for x in (edu.get("start", ""), edu.get("end", "")) if x)
            lines.append(f"{edu.get('degree', '')} — {edu.get('institution', '')} ({dates})")
        lines.append("")

    if profile.get("publications"):
        lines.append("PUBLICATIONS")
        for pub in profile["publications"]:
            suffix = " (co-authored)" if pub.get("co_authored") else ""
            meta_bits = [b for b in (pub.get("venue", ""), pub.get("date", "")) if b]
            meta = f" - {' | '.join(meta_bits)}" if meta_bits else ""
            lines.append(f"{pub.get('title', '')}{suffix}{meta}")
            if pub.get("details"):
                lines.append(f"  {pub['details']}")
        lines.append("")

    if profile.get("certifications"):
        lines += ["CERTIFICATIONS", ", ".join(profile["certifications"]), ""]

    if profile.get("languages_spoken"):
        bits = [
            f"{lang.get('language', '')} ({lang.get('level', '')})" if lang.get("level") else lang.get("language", "")
            for lang in profile["languages_spoken"]
            if lang.get("language")
        ]
        lines += ["LANGUAGES", ", ".join(bits)]

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


def render_pdf(profile: dict, out_path: Path) -> Path:
    from xml.sax.saxutils import escape

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate

    ACCENT = colors.HexColor(f"#{ACCENT_HEX}")
    RULE = colors.HexColor(f"#{RULE_HEX}")
    MUTED = colors.HexColor(f"#{MUTED_HEX}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        topMargin=0.48 * inch,
        bottomMargin=0.48 * inch,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
    )

    base = getSampleStyleSheet()["Normal"]
    name_style = ParagraphStyle("Name", parent=base, fontName="Helvetica-Bold", fontSize=22, alignment=TA_LEFT, textColor=ACCENT, spaceAfter=2)
    role_style = ParagraphStyle("Role", parent=base, fontSize=12.5, alignment=TA_LEFT, textColor=colors.HexColor("#333333"), spaceAfter=6)
    contact_style = ParagraphStyle("Contact", parent=base, fontSize=9.5, alignment=TA_LEFT, spaceAfter=6, textColor=MUTED)
    heading_style = ParagraphStyle("Heading", parent=base, fontName="Helvetica-Bold", fontSize=12, spaceBefore=13, spaceAfter=3, textColor=ACCENT)
    body_style = ParagraphStyle("Body", parent=base, fontSize=10, leading=13.5, spaceAfter=4)
    meta_style = ParagraphStyle("Meta", parent=body_style, fontName="Helvetica-Oblique", fontSize=9.5, textColor=MUTED, spaceAfter=2)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, spaceAfter=2)

    story = []

    def esc(value) -> str:
        return escape(str(value or ""))

    def add_heading(text: str) -> None:
        story.append(Paragraph(esc(text).upper(), heading_style))
        story.append(HRFlowable(width="100%", thickness=0.75, color=RULE, spaceAfter=4))

    def add_bullets(bullets: list[str]) -> None:
        items = [ListItem(Paragraph(esc(b), bullet_style), leftIndent=0) for b in bullets]
        story.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=14, bulletFontSize=9, bulletColor=ACCENT))

    contact = profile.get("contact", {})
    story.append(Paragraph(esc(contact.get("name", "").strip() or "Your Name"), name_style))

    target_roles = profile.get("target_roles") or []
    if target_roles:
        story.append(Paragraph(esc(" | ".join(target_roles[:2])), role_style))

    contact_bits = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("location", ""),
        contact.get("linkedin", ""),
        *contact.get("other_links", []),
    ]
    contact_line = " • ".join(b for b in contact_bits if b)
    if contact_line:
        story.append(Paragraph(esc(contact_line), contact_style))

    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=8))

    if profile.get("summary"):
        add_heading("Summary")
        story.append(Paragraph(esc(profile["summary"]), body_style))

    if profile.get("skills"):
        add_heading("Skills")
        story.append(Paragraph(esc(" • ".join(profile["skills"])), body_style))

    if profile.get("experience"):
        add_heading("Experience")
        for exp in profile["experience"]:
            header = f"<b>{esc(exp.get('title', ''))} — {esc(exp.get('company', ''))}</b>"
            story.append(Paragraph(header, body_style))
            dates = " – ".join(x for x in (exp.get("start", ""), exp.get("end", "")) if x)
            meta_bits = [b for b in (exp.get("location", ""), dates) if b]
            if meta_bits:
                story.append(Paragraph(esc(" | ".join(meta_bits)), meta_style))
            if exp.get("bullets"):
                add_bullets(exp["bullets"])

    if profile.get("projects"):
        add_heading("Projects")
        for proj in profile["projects"]:
            header = f"<b>{esc(proj.get('name', ''))}</b>"
            story.append(Paragraph(header, body_style))
            if proj.get("link"):
                story.append(Paragraph(esc(proj["link"]), meta_style))
            if proj.get("description"):
                story.append(Paragraph(esc(proj["description"]), body_style))
            if proj.get("bullets"):
                add_bullets(proj["bullets"])

    if profile.get("education"):
        add_heading("Education")
        for edu in profile["education"]:
            header = f"<b>{esc(edu.get('degree', ''))} — {esc(edu.get('institution', ''))}</b>"
            story.append(Paragraph(header, body_style))
            dates = " – ".join(x for x in (edu.get("start", ""), edu.get("end", "")) if x)
            meta_bits = [b for b in (edu.get("location", ""), dates) if b]
            if meta_bits:
                story.append(Paragraph(esc(" | ".join(meta_bits)), meta_style))
            if edu.get("details"):
                story.append(Paragraph(esc(edu["details"]), body_style))

    if profile.get("publications"):
        add_heading("Publications")
        for pub in profile["publications"]:
            header = f"<b>{esc(pub.get('title', ''))}</b>"
            if pub.get("co_authored"):
                header += " <i>(co-authored)</i>"
            story.append(Paragraph(header, body_style))
            meta_bits = [b for b in (pub.get("venue", ""), pub.get("date", "")) if b]
            if meta_bits:
                story.append(Paragraph(esc(" | ".join(meta_bits)), meta_style))
            if pub.get("details"):
                story.append(Paragraph(esc(pub["details"]), body_style))

    if profile.get("certifications"):
        add_heading("Certifications")
        story.append(Paragraph(esc(" • ".join(profile["certifications"])), body_style))

    if profile.get("languages_spoken"):
        add_heading("Languages")
        bits = [
            f"{lang.get('language', '')} ({lang.get('level', '')})" if lang.get("level") else lang.get("language", "")
            for lang in profile["languages_spoken"]
            if lang.get("language")
        ]
        story.append(Paragraph(esc(" • ".join(bits)), body_style))

    doc.build(story)
    return out_path
